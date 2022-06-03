import docx
from docx.shared import Inches
from docx.enum.section import WD_ORIENTATION as orient
from docx.shared import Mm
from Util import DataManager



class builder:

    images_path = []
    result_filename = ""
    images_size = []
    orientation = orient.LANDSCAPE
    layout = []

    def __init__(self, result_filename):
        self.result_filename = result_filename

    def setOrientation(self, orientations:orient):
        self.orientation = orientations

    def setLayout(self, columns, rows):
        self.layout = [columns, rows]

    def addImage(self, img_path, width, height):
        self.images_path.append(img_path)
        self.images_size.append([width, height])

    def doWrap(self):
        print("wrapping")
        doc = docx.Document()
        self.setPageSize(doc)
        self.setPageOrientation(doc)
        maxImagePerPage = self.layout[0] * self.layout[1]

        table = None
        row_cells = None
        number_of_tables = 0
        number_of_row = 0
        for i in range(len(self.images_path)):
            if i % maxImagePerPage == 0:
                table = self.getNewTable(doc)
                number_of_tables += 1
                number_of_row = 0
            if i % self.layout[0] == 0:
                row_cells = table.add_row().cells
                number_of_row += 1

            basic_cell_idx = i - ((number_of_tables-1)*maxImagePerPage)
            if i%10 == 0:
               print("done: " + str(i) + " from: " + str(len(self.images_path)))
            p = row_cells[basic_cell_idx-((number_of_row-1)*(self.layout[0]))].paragraphs[0]
            r = p.add_run()
            r.add_picture(self.images_path[i], width=Inches(self.images_size[i][0]), height=Inches(self.images_size[i][1]))
        print(str(len(self.images_path)) + " images wrapped")
        print("saving file...")
        doc.save(DataManager.RESULT_FOLDER + "/" + self.result_filename)
        print("Done!")


    def setPageOrientation(self, doc):
        section = doc.sections[-1]
        section.orientation = self.orientation

    def getNewTable(self, doc):
        return doc.add_table(rows=0, cols=self.layout[0])

    def setPageSize(self, doc):
        long = 297
        short = 210
        section = doc.sections[0]
        if self.orientation is orient.LANDSCAPE:
            section.page_height = Mm(short)
            section.page_width = Mm(long)
        if self.orientation is orient.PORTRAIT:
            section.page_height = Mm(long)
            section.page_width = Mm(short)
        section.left_margin = Mm(10)
        section.right_margin = Mm(10)
        section.top_margin = Mm(5)
        section.bottom_margin = Mm(0)
        section.header_distance = Mm(0)
        section.footer_distance = Mm(0)
